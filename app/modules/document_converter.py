"""
Semptify Document Converter Module
==================================

Converts Markdown documents to:
- Microsoft Word (.docx) with full formatting preservation
- Interactive HTML with footnotes, document links, and legal brief styling

Features:
- Markdown to DOCX conversion with headers, lists, tables, bold/italic
- Interactive HTML with numbered footnotes and citations
- Legal brief formatting with proper margins and fonts
- Document linking and cross-referencing
- Table of contents generation
- Exhibit/evidence linking
"""

import logging
import re
import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
from enum import Enum
from pydantic import BaseModel

logger = logging.getLogger(__name__)


# =============================================================================
# ENUMS & DATA MODELS
# =============================================================================

class OutputFormat(str, Enum):
    DOCX = "docx"
    HTML = "html"
    PDF = "pdf"
    BOTH = "both"


class DocumentStyle(str, Enum):
    LEGAL_BRIEF = "legal_brief"
    COURT_FILING = "court_filing"
    STANDARD = "standard"
    MEMO = "memo"


class FootnoteStyle(str, Enum):
    NUMBERED = "numbered"
    SUPERSCRIPT = "superscript"
    BRACKETED = "bracketed"


class FootnoteInfo(BaseModel):
    """Represents a footnote in the document"""
    number: int
    text: str
    citation: Optional[str] = None
    document_link: Optional[str] = None


class DocumentMetadata(BaseModel):
    """Document metadata for headers and footers"""
    title: str
    case_number: Optional[str] = None
    court: Optional[str] = None
    parties: Optional[str] = None
    author: Optional[str] = None
    date: Optional[str] = None


# =============================================================================
# MARKDOWN TO DOCX CONVERTER
# =============================================================================

class MarkdownToDocxConverter:
    """
    Converts Markdown to Microsoft Word (.docx) format
    Preserves formatting: headers, lists, tables, bold, italic, etc.
    """
    
    def __init__(self, style: DocumentStyle = DocumentStyle.STANDARD):
        self.style = style
        self.footnotes: List[FootnoteInfo] = []
        self.footnote_counter = 0
        
    def convert(self, markdown_text: str, output_path: str, 
                metadata: Optional[DocumentMetadata] = None) -> str:
        """
        Convert markdown to DOCX file
        
        Args:
            markdown_text: The markdown content to convert
            output_path: Path for the output .docx file
            metadata: Optional document metadata
            
        Returns:
            Path to the created DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX conversion. "
                "Install with: pip install python-docx"
            )
        
        doc = Document()
        self._apply_document_style(doc)
        
        # Add metadata header if provided
        if metadata:
            self._add_document_header(doc, metadata)
        
        # Parse and add content
        self._parse_markdown(doc, markdown_text)
        
        # Add footnotes section if any
        if self.footnotes:
            self._add_footnotes_section(doc)
        
        # Ensure output path has .docx extension
        if not output_path.endswith('.docx'):
            output_path += '.docx'
            
        doc.save(output_path)
        logger.info(f"Created DOCX: {output_path}")
        return output_path
    
    def _apply_document_style(self, doc):
        """Apply document-wide styling based on style type"""
        from docx.shared import Inches, Pt
        
        sections = doc.sections
        for section in sections:
            if self.style == DocumentStyle.LEGAL_BRIEF:
                # Legal brief: 1-inch margins, double-spaced
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            elif self.style == DocumentStyle.COURT_FILING:
                # Court filing: specific margins
                section.left_margin = Inches(1.5)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            else:
                # Standard document
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
    
    def _add_document_header(self, doc, metadata: DocumentMetadata):
        """Add document header with case information"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if metadata.court:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata.court.upper())
            run.bold = True
            run.font.size = Pt(12)
        
        if metadata.case_number:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Case No. {metadata.case_number}")
            run.font.size = Pt(12)
        
        if metadata.parties:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata.parties)
            run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
        
        if metadata.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata.title.upper())
            run.bold = True
            run.font.size = Pt(14)
        
        doc.add_paragraph()  # Spacing
    
    def _parse_markdown(self, doc, markdown_text: str):
        """Parse markdown and add to document"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        in_table = False
        table_rows = []
        
        while i < len(lines):
            line = lines[i]
            
            # Code blocks
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block and table_rows:
                    pass  # End of code block
                i += 1
                continue
            
            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = 'No Spacing'
                run = p.runs[0] if p.runs else p.add_run()
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                i += 1
                continue
            
            # Tables
            if '|' in line and line.strip().startswith('|'):
                table_rows.append(line)
                i += 1
                continue
            elif table_rows:
                self._add_table(doc, table_rows)
                table_rows = []
            
            # Headers
            if line.startswith('#'):
                self._add_header(doc, line)
            # Horizontal rule
            elif line.strip() in ['---', '***', '___']:
                doc.add_paragraph('_' * 50)
            # Bullet lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                self._add_list_item(doc, line, 'bullet')
            # Numbered lists
            elif re.match(r'^\s*\d+\.\s', line):
                self._add_list_item(doc, line, 'number')
            # Blockquote
            elif line.strip().startswith('>'):
                self._add_blockquote(doc, line)
            # Empty line
            elif not line.strip():
                doc.add_paragraph()
            # Regular paragraph
            else:
                self._add_paragraph(doc, line)
            
            i += 1
        
        # Handle remaining table rows
        if table_rows:
            self._add_table(doc, table_rows)
    
    def _add_header(self, doc, line: str):
        """Add a header to the document"""
        level = len(re.match(r'^#+', line).group())
        text = re.sub(r'^#+\s*', '', line)
        
        # Clean formatting from text
        text = self._strip_formatting(text)
        
        if level == 1:
            doc.add_heading(text, level=0)
        else:
            doc.add_heading(text, level=min(level, 9))
    
    def _add_paragraph(self, doc, line: str):
        """Add a formatted paragraph"""
        p = doc.add_paragraph()
        self._add_formatted_text(p, line)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with bold, italic, and footnote formatting"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Extract footnotes first
        text, footnotes = self._extract_footnotes(text)
        
        # Parse inline formatting
        pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|__.+?__|_.+?_|~~.+?~~|`[^`]+`|\[.+?\]\(.+?\))'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold and italic (***text***)
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            # Bold (**text** or __text__)
            elif (part.startswith('**') and part.endswith('**')) or \
                 (part.startswith('__') and part.endswith('__')):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # Italic (*text* or _text_)
            elif (part.startswith('*') and part.endswith('*')) or \
                 (part.startswith('_') and part.endswith('_')):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # Strikethrough (~~text~~)
            elif part.startswith('~~') and part.endswith('~~'):
                run = paragraph.add_run(part[2:-2])
                run.font.strike = True
            # Inline code (`text`)
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Links [text](url)
            elif re.match(r'\[.+?\]\(.+?\)', part):
                match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if match:
                    link_text = match.group(1)
                    # Just add the text (hyperlinks in python-docx are complex)
                    run = paragraph.add_run(link_text)
                    run.underline = True
            else:
                paragraph.add_run(part)
        
        # Add footnote references
        for fn in footnotes:
            run = paragraph.add_run(f"[{fn.number}]")
            run.font.superscript = True
    
    def _extract_footnotes(self, text: str) -> Tuple[str, List[FootnoteInfo]]:
        """Extract footnote markers and return cleaned text"""
        footnotes = []
        
        # Pattern for [^1] or [^note] style footnotes
        pattern = r'\[\^(\w+)\]'
        
        def replace_fn(match):
            self.footnote_counter += 1
            fn = FootnoteInfo(
                number=self.footnote_counter,
                text=match.group(1)
            )
            footnotes.append(fn)
            self.footnotes.append(fn)
            return ''
        
        cleaned_text = re.sub(pattern, replace_fn, text)
        return cleaned_text, footnotes
    
    def _add_list_item(self, doc, line: str, list_type: str):
        """Add a list item"""
        # Remove list marker
        if list_type == 'bullet':
            text = re.sub(r'^\s*[-*]\s+', '', line)
        else:
            text = re.sub(r'^\s*\d+\.\s+', '', line)
        
        # Determine indentation level
        indent = len(line) - len(line.lstrip())
        level = indent // 2
        
        p = doc.add_paragraph(style='List Bullet' if list_type == 'bullet' else 'List Number')
        self._add_formatted_text(p, text)
    
    def _add_blockquote(self, doc, line: str):
        """Add a blockquote"""
        from docx.shared import Inches, Pt
        
        text = re.sub(r'^>\s*', '', line)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(text)
        run.italic = True
    
    def _add_table(self, doc, table_rows: List[str]):
        """Add a table from markdown table rows"""
        if len(table_rows) < 2:
            return
        
        # Parse table
        def parse_row(row):
            cells = row.strip().strip('|').split('|')
            return [cell.strip() for cell in cells]
        
        rows_data = [parse_row(row) for row in table_rows]
        
        # Skip separator row (containing ---)
        rows_data = [r for r in rows_data if not all('-' in c for c in r)]
        
        if not rows_data:
            return
        
        num_cols = max(len(r) for r in rows_data)
        num_rows = len(rows_data)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text
        
        doc.add_paragraph()  # Spacing after table
    
    def _add_footnotes_section(self, doc):
        """Add footnotes section at the end"""
        from docx.shared import Pt
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        
        p = doc.add_paragraph()
        run = p.add_run('FOOTNOTES')
        run.bold = True
        
        for fn in self.footnotes:
            p = doc.add_paragraph()
            run = p.add_run(f"[{fn.number}] ")
            run.font.superscript = True
            p.add_run(fn.text)
            if fn.citation:
                p.add_run(f" ({fn.citation})")
    
    def _strip_formatting(self, text: str) -> str:
        """Remove markdown formatting from text"""
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text


# =============================================================================
# MARKDOWN TO INTERACTIVE HTML CONVERTER
# =============================================================================

class MarkdownToHtmlConverter:
    """
    Converts Markdown to interactive HTML with:
    - Footnotes with popup/hover display
    - Document links and cross-references
    - Legal brief styling
    - Table of contents
    - Print-friendly formatting
    """
    
    def __init__(self, style: DocumentStyle = DocumentStyle.LEGAL_BRIEF):
        self.style = style
        self.footnotes: List[FootnoteInfo] = []
        self.footnote_counter = 0
        self.document_links: Dict[str, str] = {}
        self.toc_entries: List[Dict[str, Any]] = []
    
    def convert(self, markdown_text: str, output_path: str,
                metadata: Optional[DocumentMetadata] = None,
                linked_documents: Optional[Dict[str, str]] = None) -> str:
        """
        Convert markdown to interactive HTML
        
        Args:
            markdown_text: The markdown content
            output_path: Path for output HTML file
            metadata: Document metadata
            linked_documents: Dict of document names to file paths for linking
            
        Returns:
            Path to the created HTML file
        """
        if linked_documents:
            self.document_links = linked_documents
        
        # Parse footnotes first
        markdown_text, footnote_defs = self._parse_footnote_definitions(markdown_text)
        
        # Convert markdown to HTML
        content_html = self._parse_markdown(markdown_text)
        
        # Build full HTML document
        html = self._build_html_document(content_html, metadata, footnote_defs)
        
        # Ensure output path has .html extension
        if not output_path.endswith('.html'):
            output_path += '.html'
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        logger.info(f"Created HTML: {output_path}")
        return output_path
    
    def _parse_footnote_definitions(self, text: str) -> Tuple[str, Dict[str, str]]:
        """Extract footnote definitions from markdown"""
        footnote_defs = {}
        
        # Pattern for [^1]: Footnote text
        pattern = r'^\[\^(\w+)\]:\s*(.+)$'
        
        lines = text.split('\n')
        new_lines = []
        
        for line in lines:
            match = re.match(pattern, line)
            if match:
                key = match.group(1)
                value = match.group(2)
                footnote_defs[key] = value
            else:
                new_lines.append(line)
        
        return '\n'.join(new_lines), footnote_defs
    
    def _parse_markdown(self, text: str) -> str:
        """Parse markdown and convert to HTML"""
        import html as html_lib
        
        lines = text.split('\n')
        html_parts = []
        in_code_block = False
        code_lang = ''
        code_lines = []
        in_list = False
        list_type = None
        list_items = []
        in_table = False
        table_rows = []
        
        for i, line in enumerate(lines):
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    code_content = html_lib.escape('\n'.join(code_lines))
                    html_parts.append(
                        f'<pre><code class="language-{code_lang}">{code_content}</code></pre>'
                    )
                    in_code_block = False
                    code_lines = []
                    code_lang = ''
                else:
                    # Start code block
                    in_code_block = True
                    code_lang = line.strip()[3:] or 'text'
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            # Tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    # Flush any pending list
                    if in_list:
                        html_parts.append(self._render_list(list_items, list_type))
                        in_list = False
                        list_items = []
                    in_table = True
                table_rows.append(line)
                continue
            elif in_table:
                html_parts.append(self._render_table(table_rows))
                in_table = False
                table_rows = []
            
            # Lists
            is_bullet = line.strip().startswith('- ') or line.strip().startswith('* ')
            is_number = bool(re.match(r'^\s*\d+\.\s', line))
            
            if is_bullet or is_number:
                new_list_type = 'ul' if is_bullet else 'ol'
                if not in_list:
                    in_list = True
                    list_type = new_list_type
                elif list_type != new_list_type:
                    html_parts.append(self._render_list(list_items, list_type))
                    list_items = []
                    list_type = new_list_type
                
                # Extract list item text
                if is_bullet:
                    item_text = re.sub(r'^\s*[-*]\s+', '', line)
                else:
                    item_text = re.sub(r'^\s*\d+\.\s+', '', line)
                list_items.append(item_text)
                continue
            elif in_list:
                html_parts.append(self._render_list(list_items, list_type))
                in_list = False
                list_items = []
            
            # Headers
            if line.startswith('#'):
                html_parts.append(self._render_header(line))
            # Horizontal rule
            elif line.strip() in ['---', '***', '___']:
                html_parts.append('<hr class="divider">')
            # Blockquote
            elif line.strip().startswith('>'):
                text = re.sub(r'^>\s*', '', line)
                html_parts.append(f'<blockquote>{self._format_inline(text)}</blockquote>')
            # Empty line
            elif not line.strip():
                html_parts.append('')
            # Regular paragraph
            else:
                html_parts.append(f'<p>{self._format_inline(line)}</p>')
        
        # Flush remaining content
        if in_code_block:
            code_content = html_lib.escape('\n'.join(code_lines))
            html_parts.append(f'<pre><code>{code_content}</code></pre>')
        if in_list:
            html_parts.append(self._render_list(list_items, list_type))
        if in_table:
            html_parts.append(self._render_table(table_rows))
        
        return '\n'.join(html_parts)
    
    def _render_header(self, line: str) -> str:
        """Render a header and add to TOC"""
        level = len(re.match(r'^#+', line).group())
        text = re.sub(r'^#+\s*', '', line)
        text_clean = self._strip_formatting(text)
        
        # Generate ID for linking
        header_id = re.sub(r'[^\w\s-]', '', text_clean.lower())
        header_id = re.sub(r'\s+', '-', header_id)
        
        # Add to TOC
        self.toc_entries.append({
            'level': level,
            'text': text_clean,
            'id': header_id
        })
        
        return f'<h{level} id="{header_id}">{self._format_inline(text)}</h{level}>'
    
    def _render_list(self, items: List[str], list_type: str) -> str:
        """Render a list"""
        item_html = '\n'.join(f'<li>{self._format_inline(item)}</li>' for item in items)
        return f'<{list_type}>\n{item_html}\n</{list_type}>'
    
    def _render_table(self, rows: List[str]) -> str:
        """Render a table"""
        def parse_row(row):
            cells = row.strip().strip('|').split('|')
            return [cell.strip() for cell in cells]
        
        rows_data = [parse_row(row) for row in rows]
        
        # Find separator row
        sep_idx = -1
        for i, row in enumerate(rows_data):
            if all(re.match(r'^[-:]+$', c) for c in row if c):
                sep_idx = i
                break
        
        html = '<table class="data-table">\n'
        
        if sep_idx > 0:
            # Has header
            html += '<thead>\n<tr>\n'
            for cell in rows_data[0]:
                html += f'<th>{self._format_inline(cell)}</th>\n'
            html += '</tr>\n</thead>\n'
            body_start = sep_idx + 1
        else:
            body_start = 0
        
        html += '<tbody>\n'
        for row in rows_data[body_start:]:
            html += '<tr>\n'
            for cell in row:
                html += f'<td>{self._format_inline(cell)}</td>\n'
            html += '</tr>\n'
        html += '</tbody>\n</table>'
        
        return html
    
    def _format_inline(self, text: str) -> str:
        """Format inline markdown elements"""
        import html as html_lib
        
        # Escape HTML first
        # text = html_lib.escape(text)  # Don't escape - we need to preserve some HTML
        
        # Footnotes [^1]
        def footnote_replace(match):
            key = match.group(1)
            self.footnote_counter += 1
            num = self.footnote_counter
            self.footnotes.append(FootnoteInfo(number=num, text=key))
            return f'<sup class="footnote-ref" data-footnote="{key}"><a href="#fn-{key}">[{num}]</a></sup>'
        
        text = re.sub(r'\[\^(\w+)\]', footnote_replace, text)
        
        # Bold and italic (***text***)
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', text)
        
        # Bold (**text** or __text__)
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.+?)__', r'<strong>\1</strong>', text)
        
        # Italic (*text* or _text_)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = re.sub(r'(?<!\w)_(.+?)_(?!\w)', r'<em>\1</em>', text)
        
        # Strikethrough (~~text~~)
        text = re.sub(r'~~(.+?)~~', r'<del>\1</del>', text)
        
        # Inline code (`text`)
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        
        # Links [text](url)
        def link_replace(match):
            link_text = match.group(1)
            url = match.group(2)
            # Check if it's a document link
            if url in self.document_links:
                url = self.document_links[url]
                return f'<a href="{url}" class="document-link" target="_blank">{link_text}</a>'
            return f'<a href="{url}">{link_text}</a>'
        
        text = re.sub(r'\[(.+?)\]\((.+?)\)', link_replace, text)
        
        return text
    
    def _strip_formatting(self, text: str) -> str:
        """Remove markdown formatting"""
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'\[\^(\w+)\]', '', text)
        return text
    
    def _generate_toc(self) -> str:
        """Generate table of contents HTML"""
        if not self.toc_entries:
            return ''
        
        html = '<nav class="toc">\n<h2>Table of Contents</h2>\n<ul>\n'
        
        for entry in self.toc_entries:
            indent = '  ' * (entry['level'] - 1)
            html += f'{indent}<li><a href="#{entry["id"]}">{entry["text"]}</a></li>\n'
        
        html += '</ul>\n</nav>'
        return html
    
    def _generate_footnotes_html(self, footnote_defs: Dict[str, str]) -> str:
        """Generate footnotes section"""
        if not self.footnotes and not footnote_defs:
            return ''
        
        html = '<section class="footnotes">\n<h2>Footnotes</h2>\n<ol>\n'
        
        for fn in self.footnotes:
            fn_text = footnote_defs.get(fn.text, fn.text)
            html += f'<li id="fn-{fn.text}">{fn_text} <a href="#" class="footnote-back">↩</a></li>\n'
        
        html += '</ol>\n</section>'
        return html
    
    def _build_html_document(self, content: str, metadata: Optional[DocumentMetadata],
                            footnote_defs: Dict[str, str]) -> str:
        """Build complete HTML document"""
        
        title = metadata.title if metadata else 'Document'
        
        # Generate TOC after parsing content
        toc_html = self._generate_toc()
        footnotes_html = self._generate_footnotes_html(footnote_defs)
        
        # Build header
        header_html = ''
        if metadata:
            header_html = '<header class="document-header">\n'
            if metadata.court:
                header_html += f'<div class="court">{metadata.court.upper()}</div>\n'
            if metadata.case_number:
                header_html += f'<div class="case-number">Case No. {metadata.case_number}</div>\n'
            if metadata.parties:
                header_html += f'<div class="parties">{metadata.parties}</div>\n'
            if metadata.title:
                header_html += f'<h1 class="document-title">{metadata.title}</h1>\n'
            if metadata.date:
                header_html += f'<div class="date">{metadata.date}</div>\n'
            header_html += '</header>\n'
        
        # Get CSS
        css = self._get_css()
        
        # Get JavaScript for interactivity
        js = self._get_javascript()
        
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
{css}
    </style>
</head>
<body>
    <div class="document-container">
        {header_html}
        {toc_html}
        <main class="content">
{content}
        </main>
        {footnotes_html}
    </div>
    <div id="footnote-popup" class="footnote-popup"></div>
    <script>
{js}
    </script>
</body>
</html>'''
    
    def _get_css(self) -> str:
        """Get CSS styles for the document"""
        
        base_css = '''
        :root {
            --primary-color: #1a365d;
            --secondary-color: #2c5282;
            --accent-color: #3182ce;
            --text-color: #1a202c;
            --bg-color: #ffffff;
            --border-color: #e2e8f0;
            --footnote-bg: #f7fafc;
        }
        
        * {
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.6;
            color: var(--text-color);
            background: #f5f5f5;
            margin: 0;
            padding: 20px;
        }
        
        .document-container {
            max-width: 8.5in;
            margin: 0 auto;
            background: var(--bg-color);
            padding: 1in;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }
        
        /* Header Styles */
        .document-header {
            text-align: center;
            margin-bottom: 2em;
            padding-bottom: 1em;
            border-bottom: 2px solid var(--primary-color);
        }
        
        .court {
            font-weight: bold;
            font-size: 14pt;
            margin-bottom: 0.5em;
        }
        
        .case-number {
            font-size: 12pt;
            margin-bottom: 0.5em;
        }
        
        .parties {
            font-size: 12pt;
            margin-bottom: 1em;
        }
        
        .document-title {
            font-size: 16pt;
            font-weight: bold;
            text-transform: uppercase;
            margin: 1em 0 0.5em;
        }
        
        .date {
            font-size: 11pt;
            color: #666;
        }
        
        /* Table of Contents */
        .toc {
            background: var(--footnote-bg);
            border: 1px solid var(--border-color);
            padding: 1em 1.5em;
            margin-bottom: 2em;
            border-radius: 4px;
        }
        
        .toc h2 {
            font-size: 14pt;
            margin-top: 0;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 0.5em;
        }
        
        .toc ul {
            list-style: none;
            padding-left: 0;
        }
        
        .toc li {
            margin: 0.3em 0;
        }
        
        .toc a {
            color: var(--accent-color);
            text-decoration: none;
        }
        
        .toc a:hover {
            text-decoration: underline;
        }
        
        /* Content Styles */
        h1, h2, h3, h4, h5, h6 {
            color: var(--primary-color);
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        
        h1 { font-size: 18pt; }
        h2 { font-size: 16pt; }
        h3 { font-size: 14pt; }
        h4 { font-size: 13pt; }
        h5 { font-size: 12pt; }
        h6 { font-size: 11pt; }
        
        p {
            margin: 1em 0;
            text-align: justify;
        }
        
        blockquote {
            margin: 1em 2em;
            padding: 0.5em 1em;
            border-left: 4px solid var(--accent-color);
            background: var(--footnote-bg);
            font-style: italic;
        }
        
        /* Tables */
        .data-table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }
        
        .data-table th,
        .data-table td {
            border: 1px solid var(--border-color);
            padding: 0.5em;
            text-align: left;
        }
        
        .data-table th {
            background: var(--primary-color);
            color: white;
            font-weight: bold;
        }
        
        .data-table tr:nth-child(even) {
            background: var(--footnote-bg);
        }
        
        /* Code */
        code {
            background: #f1f5f9;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 0.9em;
        }
        
        pre {
            background: #1e293b;
            color: #e2e8f0;
            padding: 1em;
            border-radius: 4px;
            overflow-x: auto;
        }
        
        pre code {
            background: none;
            padding: 0;
            color: inherit;
        }
        
        /* Footnotes */
        .footnote-ref {
            cursor: pointer;
        }
        
        .footnote-ref a {
            color: var(--accent-color);
            text-decoration: none;
        }
        
        .footnotes {
            margin-top: 3em;
            padding-top: 1em;
            border-top: 2px solid var(--border-color);
        }
        
        .footnotes h2 {
            font-size: 14pt;
        }
        
        .footnotes ol {
            padding-left: 1.5em;
        }
        
        .footnotes li {
            margin: 0.5em 0;
            font-size: 10pt;
        }
        
        .footnote-back {
            color: var(--accent-color);
            text-decoration: none;
            margin-left: 0.5em;
        }
        
        /* Footnote Popup */
        .footnote-popup {
            display: none;
            position: fixed;
            max-width: 400px;
            background: white;
            border: 1px solid var(--border-color);
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
            padding: 1em;
            border-radius: 4px;
            font-size: 10pt;
            z-index: 1000;
        }
        
        .footnote-popup.visible {
            display: block;
        }
        
        /* Document Links */
        .document-link {
            color: var(--accent-color);
            font-weight: bold;
            text-decoration: none;
            border-bottom: 1px dashed var(--accent-color);
        }
        
        .document-link:hover {
            background: #ebf8ff;
        }
        
        .document-link::before {
            content: "📄 ";
        }
        
        /* Divider */
        .divider {
            border: none;
            border-top: 1px solid var(--border-color);
            margin: 2em 0;
        }
        
        /* Print Styles */
        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .document-container {
                box-shadow: none;
                max-width: none;
                padding: 0;
            }
            
            .footnote-popup {
                display: none !important;
            }
            
            a {
                color: black;
                text-decoration: none;
            }
        }
        '''
        
        return base_css
    
    def _get_javascript(self) -> str:
        """Get JavaScript for interactivity"""
        return '''
        document.addEventListener('DOMContentLoaded', function() {
            const popup = document.getElementById('footnote-popup');
            const footnoteRefs = document.querySelectorAll('.footnote-ref');
            
            footnoteRefs.forEach(ref => {
                ref.addEventListener('mouseenter', function(e) {
                    const footnoteKey = this.dataset.footnote;
                    const footnoteEl = document.getElementById('fn-' + footnoteKey);
                    
                    if (footnoteEl) {
                        popup.innerHTML = footnoteEl.innerHTML;
                        popup.classList.add('visible');
                        
                        // Position popup
                        const rect = this.getBoundingClientRect();
                        popup.style.left = rect.left + 'px';
                        popup.style.top = (rect.bottom + 10) + 'px';
                    }
                });
                
                ref.addEventListener('mouseleave', function() {
                    popup.classList.remove('visible');
                });
            });
            
            // Smooth scroll for TOC links
            document.querySelectorAll('.toc a').forEach(link => {
                link.addEventListener('click', function(e) {
                    e.preventDefault();
                    const target = document.querySelector(this.getAttribute('href'));
                    if (target) {
                        target.scrollIntoView({ behavior: 'smooth' });
                    }
                });
            });
        });
        '''


# =============================================================================
# UNIFIED DOCUMENT CONVERTER API
# =============================================================================

class DocumentConverter:
    """
    Unified API for document conversion
    Supports Markdown to DOCX and HTML with legal formatting
    """
    
    def __init__(self, style: DocumentStyle = DocumentStyle.LEGAL_BRIEF):
        self.style = style
        self.docx_converter = MarkdownToDocxConverter(style)
        self.html_converter = MarkdownToHtmlConverter(style)
    
    def convert_to_docx(self, markdown_path: str, output_path: Optional[str] = None,
                       metadata: Optional[DocumentMetadata] = None) -> str:
        """
        Convert a markdown file to DOCX
        
        Args:
            markdown_path: Path to the markdown file
            output_path: Optional output path (defaults to same name with .docx)
            metadata: Optional document metadata
            
        Returns:
            Path to the created DOCX file
        """
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        if not output_path:
            output_path = markdown_path.rsplit('.', 1)[0] + '.docx'
        
        return self.docx_converter.convert(markdown_text, output_path, metadata)
    
    def convert_to_html(self, markdown_path: str, output_path: Optional[str] = None,
                       metadata: Optional[DocumentMetadata] = None,
                       linked_documents: Optional[Dict[str, str]] = None) -> str:
        """
        Convert a markdown file to interactive HTML
        
        Args:
            markdown_path: Path to the markdown file
            output_path: Optional output path (defaults to same name with .html)
            metadata: Optional document metadata
            linked_documents: Dict of document names to file paths
            
        Returns:
            Path to the created HTML file
        """
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        if not output_path:
            output_path = markdown_path.rsplit('.', 1)[0] + '.html'
        
        return self.html_converter.convert(markdown_text, output_path, metadata, linked_documents)
    
    def convert_text_to_docx(self, markdown_text: str, output_path: str,
                            metadata: Optional[DocumentMetadata] = None) -> str:
        """Convert markdown text directly to DOCX"""
        return self.docx_converter.convert(markdown_text, output_path, metadata)
    
    def convert_text_to_html(self, markdown_text: str, output_path: str,
                            metadata: Optional[DocumentMetadata] = None,
                            linked_documents: Optional[Dict[str, str]] = None) -> str:
        """Convert markdown text directly to HTML"""
        return self.html_converter.convert(markdown_text, output_path, metadata, linked_documents)
    
    def convert_to_both(self, markdown_path: str, output_dir: Optional[str] = None,
                       metadata: Optional[DocumentMetadata] = None,
                       linked_documents: Optional[Dict[str, str]] = None) -> Dict[str, str]:
        """
        Convert markdown to both DOCX and HTML
        
        Returns:
            Dict with 'docx' and 'html' paths
        """
        base_path = markdown_path.rsplit('.', 1)[0]
        
        if output_dir:
            filename = os.path.basename(base_path)
            base_path = os.path.join(output_dir, filename)
        
        docx_path = self.convert_to_docx(markdown_path, base_path + '.docx', metadata)
        html_path = self.convert_to_html(markdown_path, base_path + '.html', metadata, linked_documents)
        
        return {
            'docx': docx_path,
            'html': html_path
        }


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def markdown_to_docx(markdown_path: str, output_path: Optional[str] = None,
                    style: DocumentStyle = DocumentStyle.LEGAL_BRIEF,
                    case_number: Optional[str] = None,
                    title: Optional[str] = None) -> str:
    """
    Quick function to convert markdown to DOCX
    
    Example:
        markdown_to_docx('brief.md', 'brief.docx', case_number='19AV-CV-25-3477')
    """
    metadata = None
    if case_number or title:
        metadata = DocumentMetadata(
            title=title or 'Legal Document',
            case_number=case_number
        )
    
    converter = DocumentConverter(style)
    return converter.convert_to_docx(markdown_path, output_path, metadata)


def markdown_to_html(markdown_path: str, output_path: Optional[str] = None,
                    style: DocumentStyle = DocumentStyle.LEGAL_BRIEF,
                    case_number: Optional[str] = None,
                    title: Optional[str] = None,
                    linked_documents: Optional[Dict[str, str]] = None) -> str:
    """
    Quick function to convert markdown to interactive HTML
    
    Example:
        markdown_to_html('brief.md', 'brief.html', 
                        case_number='19AV-CV-25-3477',
                        linked_documents={'Exhibit A': 'exhibits/a.pdf'})
    """
    metadata = None
    if case_number or title:
        metadata = DocumentMetadata(
            title=title or 'Legal Document',
            case_number=case_number
        )
    
    converter = DocumentConverter(style)
    return converter.convert_to_html(markdown_path, output_path, metadata, linked_documents)


# =============================================================================
# MAIN - CLI USAGE
# =============================================================================

if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python document_converter.py <markdown_file> [output_format]")
        print("  output_format: docx, html, or both (default: both)")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'both'
    
    converter = DocumentConverter()
    
    if output_format == 'docx':
        result = converter.convert_to_docx(input_file)
        print(f"Created: {result}")
    elif output_format == 'html':
        result = converter.convert_to_html(input_file)
        print(f"Created: {result}")
    else:
        results = converter.convert_to_both(input_file)
        print(f"Created DOCX: {results['docx']}")
        print(f"Created HTML: {results['html']}")
